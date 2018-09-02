import sqlite3
import os

abs_path = os.getcwd()


class DB_Connector():
    def __init__(self):
        self.con = sqlite3.connect(os.path.join(abs_path, 'db', 'mds_env.db'))
        self.curse = self.con.cursor()

    def init_envsurvey(self):
        self.curse.execute("""
        create table if not exists mds_env_survey (
        ID INTEGER PRIMARY KEY AUTOINCREMENT,
        name CHAR(50),
        date date,
        companyname CHAR(50),
        username CHAR(50),
        userphone CHAR(50),
        userepl CHAR(50),
        companylocation CHAR(50),
        latlng CHAR(50),
        huanping CHAR(50),
        reportbook CHAR(50),
        reporttable CHAR(50),
        dengjitable CHAR(50),
        xianchangxiangfu CHAR(50),
        xianchangxiangfu_desc TEXT,
        yanshou CHAR(50),
        yanshou_desc TEXT,
        shengchangongyi TEXT,
        yuanliang TEXT,
        chengping TEXT,
        weixianping CHAR(50),
        wuranqingkuang CHAR(50),
        huanjingjijinyuyan CHAR(50),
        qingjiebianhao CHAR(50),
        wushuiqingkuang CHAR(50),
        dunwei FLOAT(50),
        wushuichuligongyi CHAR(50),
        feiqiqingkuang CHAR(50),
        fengliang FLOAT(50),
        feiqichuligongyi CHAR(50),
        wuni FLOAT(50),
        youqi FLOAT(50),
        jinshu FLOAT(50),
        fenchen FLOAT(50),
        feiqiwutianxiemingcheng FLOAT(50),
        qitashuoming TEXT,
        zaoyinqingkuang CHAR(50),
        pianjian FLOAT(50),
        PAC FLOAT(50),
        PAM FLOAT(50),
        tansuangai FLOAT(50),
        chulinji FLOAT(50),
        nalixianggaizao CHAR(50),
        zhushi TEXT,
        pic TEXT
        )
        """)

    def drop(self):
        self.curse.execute("drop table if exists mds_env_survey")

    def close(self):
        self.curse.close()
        self.con.close()



def download_all():
    db = DB_Connector()
    db.curse.execute("select * from mds_env_survey")
    data = db.curse.fetchall()
    dt = [i[1:-1] for i in data]
    pics = {i[-1]: i[3] for i in data if i[-1] and i[3]}
    all_cols_table = """业务人员,日期,企业名称,业主姓名,业主电话,业主职务,企业地址,经纬度,环评情况,报告书,报告表,登记表,现场相符,不否说明,验收情况,为什么没验收,生产工艺,原辅料,成品,危险品,重点污染,环境应急预案,清洁生产,污水,吨位,污水处理,废气,风量,废气处理,污泥,油漆,金属,粉尘,废弃物,其他物质说明,噪音,片碱,PAC,PAM,碳酸钙,除磷剂,哪里需要改造,注"""
    cols = all_cols_table.split(',')
    dt = [cols] + dt
    import xlwt, os, zipfile
    wbk = xlwt.Workbook()
    sheet = wbk.add_sheet('Sheet1', cell_overwrite_ok=True)
    for i in range(len(dt)):
        for j in range(len(dt[i])):
            sheet.write(i, j, dt[i][j])
    wbk.save('static/download_all/all.xls')
    for k, v in pics.items():
        if not os.path.exists(os.path.join(os.getcwd(), 'static', 'download_all', v)):
            os.mkdir(os.path.join(os.getcwd(), 'static', 'download_all', v))
        for pic in k.split(','):
            os.system("cp db/pics/%s static/download_all/%s/%s" % (pic, v, pic))
    startdir = os.path.join(os.getcwd(), 'static', 'download_all')
    file_news = os.path.join(os.getcwd(), 'static', 'download_all') + '.zip'
    z = zipfile.ZipFile(file_news, 'w', zipfile.ZIP_DEFLATED)
    for dirpath, dirnames, filenames in os.walk(startdir):
        fpath = dirpath.replace(startdir, '')
        fpath = fpath and fpath + os.sep or ''
        for filename in filenames:
            z.write(os.path.join(dirpath, filename), fpath + filename)
            print('压缩成功')
    z.close()
    os.system("rm -rf static/download_all/*")

def read_word(f_name):
    import docx
    doc_file = docx.Document(f_name)
    for i in range(len(doc_file.paragraphs)):
        print("第" + str(i) + "段的内容是：" + doc_file.paragraphs[i].text)

if __name__ == '__main__':
    read_word("requirment.docx")
    # db = DB_Connector()
    # db.curse.execute("select name,date from mds_env_survey where name in ('老王')")
    # print(db.curse.fetchall())
    # db.drop()
    # db.init_envsurvey()
